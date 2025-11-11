from pathlib import Path

from docx import Document


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value


def main() -> None:
    md_lines = Path("T2_RELATORIO.md").read_text(encoding="utf-8").splitlines()
    doc = Document()

    i = 0
    while i < len(md_lines):
        line = md_lines[i]
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("|") and line.endswith("|"):
            rows = []
            while i < len(md_lines) and md_lines[i].startswith("|"):
                row = [col.strip() for col in md_lines[i].strip("|").split("|")]
                rows.append(row)
                i += 1
            add_table(doc, rows)
            continue
        elif line.strip():
            doc.add_paragraph(line.strip())
        i += 1

    output_path = Path("T2_RELATORIO.docx")
    doc.save(output_path)
    print(f"arquivo gerado: {output_path.resolve()}")


if __name__ == "__main__":
    main()

